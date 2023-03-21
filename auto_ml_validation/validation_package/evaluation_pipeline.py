"""
Consolidate all the evaluations and generate word format report
"""

from typing import *
import pandas as pd
from validation_package.evaluation.index import ModelEvaluator
from validation_package.algorithms.logistic_regression import LogisticClassifier
from docx import Document
from docx.shared import Inches


def evaluation_pipeline(
        X_train: pd.DataFrame,
        y_train: pd.Series,
        X_test: pd.DateOffset,
        y_test: pd.Series,
        raw_train: pd.DataFrame,
        raw_test: pd.DataFrame,
        class_name_list: List[str],
        proba: pd.Series,
        model,
        benchmark_model: Optional[any],
        benchmark_proba: Optional[pd.Series],
        threshold: float
):
    evaluator = ModelEvaluator(X_train, y_train, X_test, y_test, raw_train, raw_test, class_name_list)
    model_performance = evaluator.evaluate_model(model, proba, threshold)
    if benchmark_model and benchmark_proba:
        benchmark_performance = evaluator.evaluate_model(benchmark_model, benchmark_proba, threshold)

    # generate report
    generate_report(model_performance)

def generate_report(mp, bp: Optional[any] = None):
    """
    Format and generate word report in docs.

    Input:
    - mp: model performance
    - bp: benchmark model performance. If provided, comparison will be included.

    """
    # Save images to local
    # !! for plotly chart, the package kaleido is required to export the chart
    for c in ["dist", "pr", "roc"]:
        mp[c].write_image(f'{c}.png')
    for c in ["lift"]:
        mp[c].figure.savefig(f'{c}.png')
    for c in ["pdp", "confusion"]:
        mp[c].figure_.savefig(f'{c}.png')

    doc = Document()
    doc.add_heading('Model Validation Report', 0)
    doc.add_heading('Model Performance Metrics', 1)
    doc.add_paragraph('This section includes generic model performance metrics.')
    doc.add_picture('dist.png')
    doc.add_paragraph(f"Accuracy: {mp['metrics']['accuracy']} \n Precision: {mp['metrics']['precision']} \n Recall: {mp['metrics']['recall']} \n F1 Score: {mp['metrics']['f1_score']}")
    doc.add_picture('lift.png')
    doc.save('report.docx')

def main():
    X_train = pd.read_csv("~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_train_X_processed.csv")
    y_train = pd.read_csv("~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_train_y.csv")
    X_test = pd.read_csv('~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_test_X_processed.csv')
    y_test = pd.read_csv("~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_test_y.csv")
    raw_train = pd.read_csv("~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_train.csv")
    raw_test = pd.read_csv("~/Desktop/Capstone/auto-ml-validation/data/stage_2/loanstats_test.csv")
    model = LogisticClassifier()
    model.fit(X_train, y_train)
    proba = model.predict_proba(X_test)
    print("Start evaluating model.")
    evaluation_pipeline(X_train, y_train, X_test, y_test, raw_train, raw_test, None, proba, model._model, None, None, 0.5)
    print("Evaluation done!")


if __name__ == "__main__":
    main()

