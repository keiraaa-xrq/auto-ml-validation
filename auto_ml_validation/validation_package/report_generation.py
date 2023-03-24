from docx import Document
from docx.shared import Inches

def generate_report(results):
    """
    Format and generate word report in docs.

    Input:
        - results: dictionary of model's validation result
    """

    doc = Document()

    for r in results:
    # Save images to local
    # !! for plotly chart, the package kaleido is required to export the chart
        for c in ["dist", "pr", "roc"]:
            mp[c].write_image(f'{c}.png')
        for c in ["lift"]:
            mp[c].figure.savefig(f'{c}.png')
        for c in ["pdp", "confusion"]:
            mp[c].figure_.savefig(f'{c}.png')

        doc.add_section()
        doc.add_heading('Model Validation Report', 0)
        doc.add_heading('Model Performance Metrics', 1)
        doc.add_paragraph('This section includes generic model performance metrics.')
        doc.add_picture('dist.png')
        doc.add_paragraph(f"Accuracy: {mp['metrics']['accuracy']} \n Precision: {mp['metrics']['precision']} \n Recall: {mp['metrics']['recall']} \n F1 Score: {mp['metrics']['f1_score']}")
        doc.add_picture('lift.png')
        doc.save('report.docx')