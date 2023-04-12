from typing import *
import logging
from docx import Document
from .util import *
from ..utils.logger import log_error

logger = logging.getLogger("main."+__name__)


def generate_report(
    m_result,
    bm_result: Optional[dict] = None,
    report_path: str = 'report.docx'
):

    doc = Document()
    doc.add_heading('Model Validation Report', 0)

    doc.add_heading('Model Performance Metrics', 1)
    doc.add_paragraph(
        'This section includes generic model performance metrics.')
    try:
        generate_chart_table(doc, 'dist', m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in generating distribution plot: {e}')
    try:
        for m in ['accuracy', 'precision', 'recall', 'f1_score']:
            generate_metric_table(doc, m, m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in generating metrics table: {e}')
    try:
        generate_chart_auc_table(doc, 'roc', m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in plotting ROC curve: {e}')
    try:
        generate_chart_auc_table(doc, 'pr', m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in plotting PR curve: {e}')

    doc.add_heading('Statistical Metrics', 1)
    doc.add_paragraph("This section includes statistical metrics.")
    try:
        generate_psi_table(doc, m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in generating PSI: {e}')
        raise e
    try:
        samp_res = m_result[list(m_result.keys())[0]]['txt']['csi_dict']
        if len(samp_res) > 0:
            generate_csi_table(doc, m_result)
    except Exception as e:
        log_error(logger, f'Error in generating CSI: {e}')
    try:
        generate_ks_table(doc, m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in generating KS test: ({e}')
    try:
        generate_feature_gini_table(doc, m_result, bm_result)
    except Exception as e:
        log_error(logger, f'Error in generating GINI: {e}')

    try:
        doc.add_heading('Transparency Metrics', 1)
        doc.add_paragraph(
            "This section includes LIME and SHAP interpretability under the framework of MAS's FEAT metrics.")
        doc.add_paragraph(interp_map['global_']['exp'])

        generate_trans_table(doc, 'lime', m_result, bm_result)
        generate_trans_table(doc, 'shap', m_result, bm_result)
    except:
        log_error(logger, "Error in generating transparency report")

    doc.save(report_path)
