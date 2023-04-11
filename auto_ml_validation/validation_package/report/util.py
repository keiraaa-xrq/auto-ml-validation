from typing import Dict, Optional
from docx.shared import Inches, Pt
import pathlib

# Metric Interpretation
interp_map = {
    "dist": {
        "title": "Prediction Distribution",
        "exp": "Probability distribution from model prediction output.",
    },
    "confusion": {
        "title": "Confusion Matrix",
        "exp": "Provides the correctness of the model classification on different classes. \nTrue Positive: predicted positive and it's true \nTrue Negative: predicted negative and it's true \nFalse Positive (Type 1 Error): predicted positive and it's false \nFalse Negative (Type 2 Error): predicted negative and it's false \n Please refer to the link for more details: https://towardsdatascience.com/understanding-confusion-matrix-a9ad42dcfd62)",
    },
    "roc": {
        "title": "Receiver Operating Characteristic (ROC) Curve",
        "exp": "Shows the performance of the model at all classification thresholds. Higher AUC indicates a better model.",
    },
    "pr": {
        "title": "Precision Recall Curve",
        "exp": "Shows the tradeoff between precision and recall for different thresholds in binary classification problems. A high area under the curve represents both high recall and high precision, where high precision relates to a low false positive rate, and high recall relates to a low false negative rate. Please refer to the link for more details: https://medium.com/@douglaspsteen/precision-recall-curves-d32e5b290248",
    },
    "lift": {
        "title": "Lift Curve",
        "exp": "The ratio of the number of positive observations up to decile i using the model to the expected number of positives up to that decile is based on a random model. Lift chart is the chart between the lift on the vertical axis and the corresponding decile on the horizontal axis. The Greater the area between the Lift / Gain and Baseline, the Better the model. Please refer to the link for more details: https://www.geeksforgeeks.org/understanding-gain-chart-and-lift-chart/",
    },
    "accuracy": {
        "title": "Accuracy",
        "exp": "From all the classes, how many are predicted correctly.",
    },
    "precision": {
        "title": "Precision",
        "exp": "= True Positive / (True Positive + False Positive) \n From all that are predicted as positive, how many are actually positive.",
    },
    "recall": {
        "title": "Recall",
        "exp": "= True Positive / (True Positive + False Negative) \n From all the positive classes, how many are predicted correctly.",
    },
    "f1_score": {
        "title": "F1 Score",
        "exp": "= 2*Recall*Precision / (Recall + Precision). \n F1-score helps to measure Recall and Precision at the same time.",
    },
    "feature_gini": {
        "title": "Gini Impurity",
        "exp": "A measurement used to build Decision Trees to determine how the features of a dataset should split nodes to form the tree. It calculates the probability of a certain randomly selected feature that is classified incorrectly. An attribute with the smallest Gini Impurity should be selected for splitting the node. Please refer to the link for more details: https://www.learndatasci.com/glossary/gini-impurity/",
    },
    "dataset_gini": {
        "title": "Normalized Gini Index",
        "exp": "An adjusted way to represent AUC. A perfectly random model will have a value of 0, the reversing model having a negative sign and the perfect model having 1. Please refer to the link for more details: https://yassineelkhal.medium.com/confusion-matrix-auc-and-roc-curve-and-gini-clearly-explained-221788618eb2",
    },
    "lime": {
        "title": "LIME Interpretability",
        "exp": "LIME is a method that helps to explain the predictions of any machine learning model by approximating the model locally with an interpretable model. LIME works by perturbing the input data and observing the changes in the output of the model. The perturbed data is then used to train a new, interpretable model that approximates the behavior of the original model in that local region. The resulting model can then be used to explain why the original model made a particular prediction. \n The LIME chart shows the coefficients assigned to each feature by the interpretable model. The coefficients indicate the direction and magnitude of the impact of each feature on the model's output. Positive coefficients indicate that the feature has a positive impact on the Class1, while negative coefficients indicate a positive impact on Class 0. \n Please refer to the link for more details: https://arxiv.org/abs/1602.04938)",
    },
    "shap": {
        "title": "Importance",
        "exp": "SHAP is a method for interpreting the output of any machine learning model by computing the contribution of each feature to the final prediction. SHAP is based on the concept of Shapley values from cooperative game theory, which measures the marginal contribution of each feature to the final prediction by averaging over all possible feature combinations. By computing the SHAP values for each feature, we can determine which features had the most significant impact on the prediction. \n The SHAP plot shows the average impact of each feature on the model's output across all instances in the dataset. The features are sorted in descending order of importance, and the magnitude and direction of the impact are shown. \n Please refer to the link for more details: https://papers.nips.cc/paper/7062-a-unified-approach-to-interpreting-model-predictions.pdf",
    },
    "global_": {
        "title": "Global",
        "exp": "Global interpretability refers to the ability to explain the behavior of a model across its entire input space. It focuses on understanding how the model works in general, rather than on a specific instance. Global interpretability techniques aim to provide insights into the relationships between features and how they contribute to the overall behavior of the model.",
    },
    "local_": {
        "title": "Local",
        "exp": "Local interpretability refers to the ability to explain the behavior of a model on a specific instance or observation. It focuses on understanding why a model made a particular prediction or decision for a single data point. Local interpretability techniques provide explanations that are specific to a particular instance, and they typically focus on the features that the model used to make its decision.",
    },
    "psi": {"title": "PSI", "exp": "A metric to measure how much a variable has shifted in distribution between two samples over time. Please refer the link for more details: https://towardsdatascience.com/psi-and-csi-top-2-model-monitoring-metrics-924a2540bed8"},
    "csi": {"title": "CSI", "exp": "It compares the distribution of an independent variable in the training data set to a testing data set. It detects shifts in the distributions of input variables that are submitted for scoring over time. Please refer the link for more details: https://towardsdatascience.com/psi-and-csi-top-2-model-monitoring-metrics-924a2540bed8"},
    "ks": {"title": "Kolmogorov–Smirnov statistic", "exp": ":Quantifies a distance of the distribution within the training sample and testing sample, or between the two. Please refer the link for more details: https://en.wikipedia.org/wiki/Kolmogorov%E2%80%93Smirnov_test"},
}


def generate_sub_heading(doc, title: str, exp: str):
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    if exp:
        doc.add_paragraph(exp)


def setup_table(
    table, title: Optional[bool] = False, include_bm: Optional[bool] = True
):
    table.rows[0].cells[0].text = "Testing Dataset"
    table.style = "TableGrid"
    table.columns[0].width = Inches(1.2)
    if title:
        table.rows[0].cells[1].text = "Model"
        if include_bm:
            table.rows[0].cells[2].text = "Benchmark Model"


def add_row(table, row_count):
    table.add_row()
    row_count += 1
    return row_count, table.rows[row_count].cells


def generate_chart_table(doc, chart_name: str, results, bm_results):
    generate_sub_heading(
        doc, interp_map.get(chart_name)[
            "title"], interp_map.get(chart_name)["exp"]
    )

    table = doc.add_table(rows=1, cols=3)
    setup_table(table)
    table.columns[2].width = Inches(6.5)
    row_count = 0
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text, row[1].text = data, "Model"
        row[2].paragraphs[0].add_run().add_picture(
            r['charts'][chart_name]
        )
        if bm_results:
            row_count, row = add_row(table, row_count)
            row[1].text = "Benchmark Model"
            row[2].paragraphs[0].add_run().add_picture(
                bm_results[data]['charts'][chart_name]
            )

    doc.add_paragraph()


def generate_chart_auc_table(doc, chart_name: str, results, bm_results):
    generate_sub_heading(
        doc, interp_map.get(chart_name)[
            "title"], interp_map.get(chart_name)["exp"]
    )

    auc_name = "ROCAUC" if chart_name == "roc" else "PRAUC"

    table = doc.add_table(rows=1, cols=4)
    setup_table(table)
    row_count = 0
    table.rows[0].cells[3].text = "AUC"
    table.columns[2].width = Inches(6.5)
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text, row[1].text = data, "Model"
        row[2].paragraphs[0].add_run().add_picture(
            r['charts'][chart_name]
        )
        row[3].text = str(round(r["txt"]["auc"][auc_name], 6))

        if bm_results:
            row_count, row = add_row(table, row_count)
            row[1].text = "Benchmark Model"
            row[2].paragraphs[0].add_run().add_picture(
                bm_results[data]['charts'][chart_name]
            )
            row[3].text = str(
                round(bm_results[data]["txt"]["auc"][auc_name], 6))

    doc.add_paragraph()


def generate_metric_table(doc, metric_name: str, results, bm_results):
    generate_sub_heading(
        doc, interp_map.get(metric_name)[
            "title"], interp_map.get(metric_name)["exp"]
    )
    cols = 3 if bm_results else 2

    table = doc.add_table(rows=1, cols=cols)
    setup_table(table, True, bm_results != None)
    row_count = 0
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text = data
        row[1].text = str(r["txt"]["metrics"].get(metric_name))
        if bm_results:
            row[2].text = str(bm_results.get(data)["txt"]
                              ["metrics"].get(metric_name))

    doc.add_paragraph()


def fill_stats_table(table, df, row_count):
    index = df.index.to_list()
    cols = df.columns
    nrow, ncol = df.shape
    row = table.rows[row_count].cells

    for i in range(nrow):
        # add index
        row[3].text = index[i]
        row[3].paragraphs[0].runs[0].font.size = Pt(8)
        # add values
        for j in range(ncol):
            row[4 + j].text = str(round(df.iloc[i, j], 6))
            # change font size
            row[4 + j].paragraphs[0].runs[0].font.size = Pt(8)
        row_count, row = add_row(table, row_count)

    return row_count


def generate_psi_table(doc, results, bm_results):
    generate_sub_heading(
        doc, interp_map.get("psi")["title"], interp_map.get("psi")["exp"]
    )

    table = doc.add_table(rows=1, cols=9)
    setup_table(table)
    header = table.rows[0].cells
    header[0].text = "Testing Dataset"
    header[2].text = "PSI"
    cols = ["train_count", "train_perc",
            "test_count", "test_perc", "index_value"]
    for i in range(5):
        header[i + 4].text = cols[i]

    row_count = 0
    row_count, row = add_row(table, row_count)
    for data, r in results.items():
        row[0].text = data
        row[1].text = "Model"
        row[2].text = str(round(r["txt"]["psi"], 8))
        row_count = fill_stats_table(table, r["txt"]["psi_df"], row_count)

        if bm_results:
            row = table.rows[row_count].cells
            row[1].text = "Benchmark Model"
            row[2].text = str(round(bm_results.get(data)["txt"]["psi"], 8))
            row_count = fill_stats_table(
                table, bm_results.get(data)["txt"]["psi_df"], row_count
            )

        row = table.rows[row_count].cells


def generate_csi_table(doc, results):
    generate_sub_heading(
        doc, interp_map.get("csi")["title"], interp_map.get("csi")["exp"]
    )

    i = 0
    for f in next(iter(results.items()))[1]['txt']['csi_dict'].keys():
        generate_sub_heading(doc, f, None)
        table = doc.add_table(rows=1, cols=9)
        setup_table(table)
        header = table.rows[0].cells
        header[0].text = "Testing Dataset"
        header[2].text = "CSI"
        cols = ["train_count", "train_perc",
                "test_count", "test_perc", "index_value"]
        for j in range(5):
            header[j + 4].text = cols[j]

        row_count = 0
        row_count, row = add_row(table, row_count)
        for data, r in results.items():
            row[0].text = data
            row[1].text = "Model"
            row[2].text = str(round(r["txt"]["csi_dict"].get(f), 8))
            row_count = fill_stats_table(
                table, r["txt"]["csi_list"][i], row_count)
            row = table.rows[row_count].cells
        i += 1


def generate_ks_table(doc, results, bm_results):
    generate_sub_heading(
        doc, interp_map.get("ks")["title"], interp_map.get("ks")["exp"]
    )

    table = doc.add_table(rows=1, cols=5)
    setup_table(table)
    header = table.rows[0].cells
    header[0].text = "Testing Dataset"
    header[2].text = "Train"
    header[3].text = "Test"
    header[4].text = "Train vs. Test"

    row_count = 0
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text = data
        row[1].text = "Model"
        row[2].text = str(round(r["txt"]["ks"]["Train"], 8))
        row[3].text = str(round(r["txt"]["ks"]["Test"], 8))
        row[4].text = str(round(r["txt"]["ks"]["Train vs Test"], 8))

        if bm_results:
            row_count, row = add_row(table, row_count)
            row[1].text = "Benchmark Model"
            row[2].text = str(
                round(bm_results.get(data)["txt"]["ks"]["Train"], 8))
            row[3].text = str(
                round(bm_results.get(data)["txt"]["ks"]["Test"], 8))
            row[4].text = str(
                round(bm_results.get(data)["txt"]["ks"]["Train vs Test"], 8)
            )


def generate_trans_table(doc, name: str, results, bm_results: Optional[any] = None):
    def _fill_table(table, row_count: int, m, bm: Optional[any] = None):
        for i in range(len(m)):
            row_count, cells = add_row(table, row_count)
            cells[1].text = m[i][0]
            cells[2].text = str(round(m[i][1], 6))
            if bm:
                cells[3].text = bm[i][0]
                cells[4].text = str(round(bm[i][1], 6))
        return row_count

    generate_sub_heading(
        doc, interp_map.get(name).get("title"), interp_map.get(name).get("exp")
    )
    cols = 5 if bm_results else 3

    for scope in ["global_"]:
        metric = scope + name
        generate_sub_heading(
            doc, f"{interp_map.get(scope)['title']} {interp_map.get(name)['title']}", None
        )
        table = doc.add_table(rows=1, cols=cols)
        setup_table(table)
        header = table.rows[0].cells
        header[1].merge(header[2])
        header[1].text = "Model"
        if bm_results:
            header[3].merge(header[4])
            header[3].text = "Benchmark Model"

        row_count = 0
        for data, r in results.items():
            row_count, row = add_row(table, row_count)
            row[0].text = data
            row[1].merge(row[2])
            row[1].paragraphs[0].add_run().add_picture(
                results[data]['charts'][metric]
            )
            if bm_results:
                row[3].merge(row[4])
                row[3].paragraphs[0].add_run().add_picture(
                    bm_results[data]['charts'][metric]
                )
                row_count = _fill_table(
                    table, row_count, r["txt"][metric], bm_results[data]["txt"][metric]
                )
            else:
                row_count = _fill_table(table, row_count, r["txt"][metric])

            row_count, row = add_row(table, row_count)

        doc.add_paragraph()


def generate_feature_gini_table(doc, results, bm_results):
    def _fill_table(table, row_count: int, m_gini: Dict[str, float], bm_gini: Optional[Dict[str, float]] = None):
        for f, v in m_gini.items():
            row_count, cells = add_row(table, row_count)
            cells[1].text = f
            cells[2].text = str(round(v, 6))
            if bm_gini:
                cells[3].text = str(round(bm_gini.get(f), 6))
        return row_count

    generate_sub_heading(doc, interp_map.get('feature_gini')[
                         'title'], interp_map.get('feature_gini')['exp'])
    cols = 4 if bm_results else 3

    table = doc.add_table(rows=1, cols=cols)
    row = table.rows[0].cells
    row[0].text = "Testing Dataset"
    row[1].text = "Feature"
    row[2].text = "Model"
    if bm_results:
        row[3].text = "Benchmark Model"
    table.style = 'TableGrid'
    table.columns[1].width = Inches(3)
    for i in range(2, cols):
        table.columns[i].width = Inches(1)

    row_count = 0
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text = data
        if bm_results:
            row_count = _fill_table(
                table, row_count, r['txt']['feature_gini'], bm_results[data]['txt']['feature_gini'])
        else:
            row_count = _fill_table(table, row_count, r['txt']['feature_gini'])

        row_count, row = add_row(table, row_count)

    doc.add_paragraph()


def generate_gini_table(doc, results, bm_results):
    generate_sub_heading(doc, interp_map.get('dataset_gini')[
                         'title'], interp_map.get('dataset_gini')['exp'])
    cols = 3 if bm_results else 2

    table = doc.add_table(rows=1, cols=cols)
    setup_table(table, True, bm_results != None)
    row_count = 0
    for data, r in results.items():
        row_count, row = add_row(table, row_count)
        row[0].text = data
        row[1].text = str(round(r['txt']['dataset_gini'], 8))
        if bm_results:
            row[2].text = str(
                round(bm_results.get(data)['txt']['dataset_gini'], 8))
    doc.add_paragraph()
